#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
GUIDE = ROOT / "research_guide"
MARKDOWN = GUIDE / "README_CN.md"
OUTPUT = GUIDE / "CropGenome-FM_详细研究设计与评估报告_CN.docx"

NAVY = "183153"
BLUE = "2F6BFF"
LIGHT_BLUE = "EAF1FF"
PALE_GREEN = "E7F4ED"
PALE_ORANGE = "FFF0E3"
GRAY = "667085"
LIGHT_GRAY = "F4F6F8"
BORDER = "C7D0DB"
WHITE = "FFFFFF"
BLACK = "20252B"

PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
    "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
    "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN",
    "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind", "contextualSpacing",
    "mirrorIndents", "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange",
]
TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd", "noWrap",
    "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark", "headers", "cellIns",
    "cellDel", "cellMerge", "tcPrChange",
]
TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
    "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription",
    "tblPrChange",
]


def insert_ordered(parent, child, order: list[str]) -> None:
    local = child.tag.rsplit("}", 1)[-1]
    target_rank = order.index(local)
    for index, existing in enumerate(parent):
        existing_local = existing.tag.rsplit("}", 1)[-1]
        if existing_local in order and order.index(existing_local) > target_rank:
            parent.insert(index, child)
            return
    parent.append(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        insert_ordered(tc_pr, shd, TCPR_ORDER)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        insert_ordered(tc_pr, tc_mar, TCPR_ORDER)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_table_borders(table, color=BORDER, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        insert_ordered(tbl_pr, borders, TBLPR_ORDER)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_run_font(run, name="Arial", size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill: str, border_color: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        insert_ordered(p_pr, borders, PPR_ORDER)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    insert_ordered(p_pr, shd, PPR_ORDER)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_toc(paragraph) -> None:
    def field_run(child):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field_run(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    field_run(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    field_run(separate)
    placeholder = OxmlElement("w:t")
    placeholder.text = "在Word中打开后，右键更新目录字段"
    field_run(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run(end)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.extend([fonts, color, underline])
    new_run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    new_run.append(node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, default_size=10.5) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()].replace("\\|", "|"))
            set_run_font(run, size=default_size, color=BLACK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True, color=BLACK)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(7.2, default_size - 0.6), color=NAVY)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:].replace("\\|", "|"))
        set_run_font(run, size=default_size, color=BLACK)


def split_table_row(line: str) -> list[str]:
    body = line.strip()[1:-1]
    return [part.strip().replace("\\|", "|").replace("<br>", "\n") for part in re.split(r"(?<!\\)\|", body)]


def clean_cell_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", lambda m: f"{m.group(1)}\n{m.group(2)}", text)
    text = text.replace("**", "").replace("`", "")
    return text


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(4)

    heading_specs = {
        "Title": (27, NAVY, 0, 12),
        "Subtitle": (18, GRAY, 0, 8),
        "Heading 1": (19, NAVY, 14, 8),
        "Heading 2": (15, NAVY, 11, 5),
        "Heading 3": (12.5, BLUE, 8, 3),
        "Heading 4": (11.2, NAVY, 6, 2),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(8.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(GRAY)

    if "Callout" not in [s.name for s in doc.styles]:
        callout = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        callout.font.name = "Arial"
        callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        callout.font.size = Pt(10.5)
        callout.font.color.rgb = RGBColor.from_string(NAVY)
        callout.paragraph_format.left_indent = Cm(0.35)
        callout.paragraph_format.right_indent = Cm(0.25)
        callout.paragraph_format.space_before = Pt(5)
        callout.paragraph_format.space_after = Pt(6)

    # Update TOC and other fields when opened in Word.
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        compat = settings.find(qn("w:compat"))
        if compat is None:
            settings.append(update)
        else:
            settings.insert(settings.index(compat), update)


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CropGenome-FM · 详细研究设计与评估报告")
    set_run_font(run, size=8, color=GRAY)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D9E1EA")
    borders.append(bottom)
    insert_ordered(p_pr, borders, PPR_ORDER)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("机器证据截止：2026-07-21  ·  文档更新：2026-07-29  ·  第 ")
    set_run_font(run, size=8, color=GRAY)
    add_page_number(fp)
    run = fp.add_run(" 页")
    set_run_font(run, size=8, color=GRAY)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(34)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "26")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    insert_ordered(p_pr, borders, PPR_ORDER)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("CropGenome-FM\n作物基因组基础模型")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("详细研究设计、数据字典、模型架构、预训练与下游评估报告")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("CURRENT EVIDENCE + PREREGISTERED NEXT-STAGE DESIGN")
    set_run_font(r, size=9, bold=True, color=WHITE)
    shade_paragraph(p, BLUE)

    metadata = [
        ("文档版本", "2.3（Stage B数据生成与区域训练覆盖完整版）"),
        ("机器证据截止", "2026-07-21 14:04 CST（UTC+08:00）"),
        ("本次文档更新", "2026-07-29（未新增或重跑formal test）"),
        ("当前冻结基座", "CropGenomeFM_step14000"),
        ("实际模型参数", "369,505,287"),
        ("交付格式", "Markdown + DOCX + 汇总数据表 + PNG/PDF图片"),
    ]
    for key, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{key}：")
        set_run_font(r, size=10.5, bold=True, color=NAVY)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=BLACK)

    p = doc.add_paragraph(style="Callout")
    p.paragraph_format.space_before = Pt(22)
    add_inline(p, "**证据边界：** 已完成8K阶段和10类正式下游评估。64K只训练到第569次参数更新，没有验证成绩；128K和256K只有DNA片段文件。NT-v2 500M和9项新任务尚未评估。", 10.8)
    shade_paragraph(p, LIGHT_BLUE, BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("CropGenome-FM Project · Reproducible Research Delivery")
    set_run_font(r, size=10, bold=True, color=GRAY)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_executive_page(doc: Document) -> None:
    doc.add_heading("执行摘要", level=1)
    bullets = [
        "已经完成：训练了一个3.695亿参数的作物DNA模型；训练在第17,000次更新停止，验证集选出第14,000次保存的模型；核心3任务和外部7任务已经正式评估。",
        "还没完成：64K只训练到第569次更新，尚未做第一次验证；128K和256K只准备了DNA片段文件，没有对应模型。",
        "当前优势：在2,048和8,192 bp核心三任务平均成绩上排名第1，也明显超过结构相同但没有预训练的对照模型。",
        "当前差距：在外部表达和lncRNA任务上，PlantCAD2和PlantCaduceus总体更强；NT-v2 500M还没有运行。",
        "指标核对：旧AUPRC代码没有正确处理大量并列分数，影响部分简单基线；用标准方法重算后，主要学习模型的排名不变。",
        "数据与抽样：258个原始版本预设为训练192、验证35、测试候选31，筛选后238个版本实际贡献Stage B片段。7类区域几乎肯定都被大量抽到，但旧日志没有逐区域真实计数；单条训练片段相对抽样权重约为理想值的0.980至1.067倍。",
        "物种边界：核心下游训练和测试不共享物种或属，但仍可能属于同一个科；3个测试基因组版本都没用于预训练，但只有黄瓜这个测试物种在预训练中完全未见。",
        "未来任务：已设计完整长基因、远端调控、多倍体、TE/SV和NLR基因簇等9项作物任务；它们目前只是实验方案，没有成绩。",
        "下一步：先建立两个最优先的长序列任务并加入NT-v2 500M，再决定是否值得投入64K–256K混合长度正式训练。",
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item, 10.5)

    image = GUIDE / "figures" / "figure_01_architecture.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Inches(6.55))
    cap = doc.add_paragraph("图1｜CropGenome-FM当前实现架构、RC双向路径与三个输出用途", style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_contents(doc: Document) -> None:
    doc.add_heading("目录｜章节导航", level=1)
    headings = []
    for line in MARKDOWN.read_text(encoding="utf-8").splitlines():
        if re.match(r"^##\s+\d+\.\s+", line) or line.startswith("# 附录"):
            headings.append(line.lstrip("# ").strip())
    midpoint = (len(headings) + 1) // 2
    left = headings[:midpoint]
    right = headings[midpoint:]
    table = doc.add_table(rows=max(len(left), len(right)), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    content_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin
    for column in table.columns:
        column.width = int(content_width / 2)
    for ridx, row in enumerate(table.rows):
        prevent_row_split(row)
        for cidx, cell in enumerate(row.cells):
            items = left if cidx == 0 else right
            text = items[ridx] if ridx < len(items) else ""
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
            set_cell_shading(cell, LIGHT_GRAY if ridx % 2 == 0 else WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=9.2, bold=text.startswith(("1.", "14.", "25.", "附录")), color=NAVY if text else WHITE)
    note = doc.add_paragraph(style="Callout")
    add_inline(note, "本页为静态章节导航，确保LibreOffice、Microsoft Word和在线预览均可直接阅读；实际页码请以页脚为准。", 9.5)
    shade_paragraph(note, LIGHT_GRAY, BORDER)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_section_orientation(section, landscape: bool) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.35)
        section.bottom_margin = Cm(1.35)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.75)
        section.bottom_margin = Cm(1.65)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True


def set_orientation(doc: Document, landscape: bool, start_type=WD_SECTION.NEW_PAGE) -> None:
    section = doc.add_section(start_type)
    configure_section_orientation(section, landscape)


def add_markdown_table(doc: Document, rows: list[list[str]], landscape_prepared: bool = False) -> bool:
    columns = len(rows[0])
    wide = columns >= 8
    if wide and not landscape_prepared:
        set_orientation(doc, True)
    section = doc.sections[-1]
    content_width = section.page_width - section.left_margin - section.right_margin
    table = doc.add_table(rows=0, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    # Widths reflect actual content while capping long URL/description columns.
    weights = []
    for col in range(columns):
        max_len = max(len(clean_cell_markdown(row[col]).replace("\n", " ")) for row in rows)
        weights.append(max(7, min(max_len, 38)))
    total_weight = sum(weights)
    widths = [int(content_width * weight / total_weight) for weight in weights]
    widths[-1] += int(content_width) - sum(widths)
    tbl_width = table._tbl.tblPr.find(qn("w:tblW"))
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(Emu(int(content_width)).twips))
    for column, width in zip(table.columns, widths):
        column.width = width
    font_size = 8.0 if columns <= 6 else (7.0 if columns <= 8 else 6.2)

    for ridx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for cidx, value in enumerate(values):
            cell = row.cells[cidx]
            cell.width = widths[cidx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ridx == 0:
                set_cell_shading(cell, NAVY)
            elif ridx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            else:
                set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            text = clean_cell_markdown(value)
            for line_index, line in enumerate(text.split("\n")):
                if line_index:
                    p.add_run().add_break()
                run = p.add_run(line)
                set_run_font(run, size=font_size, bold=(ridx == 0), color=WHITE if ridx == 0 else BLACK)
        if ridx == 0:
            set_repeat_header(row)
    if wide:
        set_orientation(doc, False, WD_SECTION.CONTINUOUS)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return wide


def add_figure(doc: Document, alt: str, target: str, number: int) -> None:
    path = (GUIDE / target).resolve()
    wide = path.stem == "figure_05_strategy_map"
    if wide:
        set_orientation(doc, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(10.0 if wide else 6.55))
    # Add descriptive metadata to the image drawing.
    for drawing in run._r.findall(qn("w:drawing")):
        for doc_pr in drawing.iter(qn("wp:docPr")):
            doc_pr.set("name", f"Figure {number}")
            doc_pr.set("title", alt)
            doc_pr.set("descr", alt)
    cap = doc.add_paragraph(f"图{number}｜{alt}", style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if wide:
        set_orientation(doc, False)


def add_markdown_body(doc: Document) -> None:
    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1. "))
    lines = lines[start:]
    i = 0
    figure_number = 1  # Architecture is already figure 1 on executive page.
    landscape_prepared = False
    just_started_portrait_section = False
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        image_match = re.fullmatch(r"!\[([^]]+)\]\(([^)]+)\)", line.strip())
        if image_match:
            figure_number += 1
            add_figure(doc, image_match.group(1), image_match.group(2), figure_number)
            just_started_portrait_section = False
            i += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = heading.group(2)
            next_index = i + 1
            while next_index < len(lines) and not lines[next_index].strip():
                next_index += 1
            next_is_wide_table = False
            if next_index < len(lines) and lines[next_index].startswith("| "):
                next_is_wide_table = len(split_table_row(lines[next_index])) >= 8
            if next_is_wide_table:
                if just_started_portrait_section:
                    configure_section_orientation(doc.sections[-1], True)
                else:
                    set_orientation(doc, True)
                landscape_prepared = True
            elif level == 1 and not just_started_portrait_section and not text.startswith("附录C"):
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            doc.add_heading(text, level=level)
            just_started_portrait_section = False
            i += 1
            continue
        if line.startswith("| "):
            block = []
            while i < len(lines) and lines[i].startswith("| "):
                block.append(split_table_row(lines[i]))
                i += 1
            if len(block) >= 2 and all(re.fullmatch(r":?-+:?", cell.replace(" ", "")) for cell in block[1]):
                block.pop(1)
            just_started_portrait_section = add_markdown_table(doc, block, landscape_prepared)
            landscape_prepared = False
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="Callout")
            add_inline(p, line[2:], 10.3)
            shade_paragraph(p, LIGHT_BLUE, BLUE)
            just_started_portrait_section = False
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, bullet.group(1), 10.2)
            just_started_portrait_section = False
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            # Keep the source number instead of Word's shared List Number
            # counter, which otherwise continues across unrelated sections.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.63)
            p.paragraph_format.first_line_indent = Cm(-0.50)
            p.paragraph_format.space_after = Pt(2)
            prefix = p.add_run(f"{numbered.group(1)}. ")
            set_run_font(prefix, size=10.2, color=BLACK)
            add_inline(p, numbered.group(2), 10.2)
            just_started_portrait_section = False
            i += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_inline(p, line, 10.5)
        just_started_portrait_section = False
        i += 1


def main() -> None:
    doc = Document()
    configure_document(doc)
    configure_header_footer(doc.sections[0])
    props = doc.core_properties
    props.title = "CropGenome-FM作物基因组基础模型：详细研究设计与评估报告"
    props.subject = "模型架构、预训练数据、正式下游结果、基线比较与下一阶段预注册设计"
    props.author = "CropGenome-FM Project"
    props.keywords = "crop genome foundation model, long context, plant genomics, benchmark"
    props.comments = "Version 2.3 Stage B data-generation and region-training coverage edition; generated from research_guide/README_CN.md and verified source-data; no new formal test was run for this documentation update."

    add_cover(doc)
    add_executive_page(doc)
    add_contents(doc)
    add_markdown_body(doc)
    doc.save(OUTPUT)
    print(f"wrote {OUTPUT} ({OUTPUT.stat().st_size} bytes, {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(doc.sections)} sections)")


if __name__ == "__main__":
    main()
